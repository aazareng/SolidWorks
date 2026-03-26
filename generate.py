"""
generate.py v2 — Polypack Manual Generation Script
===================================================
Usage:
    python generate.py --config "4628 - Pinturas Berel - Clearprint-24L.xlsx"

    # Dump extracted data as JSON (no Word doc generated — useful for debugging
    # and as a future web-API hook):
    python generate.py --config <xlsx> --json

Requirements:
    pip install python-docx openpyxl Pillow pywin32

    pywin32 + Excel (Windows) required only for image capture
    (PHYSICAL, LUBRICATION, and SCREENS screenshots).
    All other sections work without Excel installed.

Architecture (web-ready):
    load_data(xlsx)          → JSON-serializable dict   ← future API endpoint
    capture_images(xlsx, data, out_dir) → {key: Path}   ← COM, Windows-only
    build_document(data, images, template, out_path)    ← produces .docx

Document sections:
    1. Cover          ← SUMMARY machine info (text substitution)
    2. Modules        ← MOD_1…MOD_N BOM tables
    3. PM             ← PM sheet, USE=True rows, cols B:E
    4. Faults         ← FAULTS sheet, cols A:C
    5. HMI Params     ← HMI sheet (full table)
    6. HMI Screens    ← SCREENS sheet (COM-captured images + option tables)
    7. Setpoints      ← PHYSICAL sheet (single COM image)
    8. Lubrication    ← LUBRICATION sheet (single COM image)
    9. VFD Params     ← sheet matching HMI brand in SUMMARY

Future web-API notes:
    - load_data() returns a plain dict; call json.dumps(data) to serve it.
    - build_document() accepts that dict, so the web layer only needs to:
        1. Receive/store the Excel file (or accept the dict directly)
        2. Call load_data() → optionally edit fields → build_document()
    - capture_images() is the only Windows-specific piece; a future cloud
      deployment could replace it with a LibreOffice or headless-Chrome renderer.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import sys
import tempfile
from pathlib import Path

import openpyxl
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches as DInches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    import win32com.client as win32
    _WIN32_AVAILABLE = True
except ImportError:
    _WIN32_AVAILABLE = False

# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────

TEMPLATE_DOCX = Path("template.docx")
OUTPUT_DIR    = Path("output")

# SCREENS sheet layout
SCREENS_NAME_ROW = 1    # screen names
SCREENS_IMG_ROW  = 2    # "Place in Cell" images
SCREENS_HDR_ROW  = 3    # OPTION / DESCRIPTION headers
SCREENS_DATA_ROW = 4    # option rows start here

# Map HMI brand (upper-cased) → VFD sheet name
VFD_SHEET_MAP: dict[str, str] = {
    "SCHNEIDER": "SCHNEIDER",
    "AB":        "AB",
    "ALLEN-BRADLEY": "AB",
    "POWERFLEX": "POWERFLEX",
}

# ─────────────────────────────────────────────────────────────────────────────
# DATA LOADING
# Returns a JSON-serializable dict — the single source for both document
# building and any future REST/web-API layer.
# ─────────────────────────────────────────────────────────────────────────────

def load_data(xlsx_path: Path) -> dict:
    """
    Load all manual content from the machine Excel workbook.

    Returns a JSON-serializable dict with keys:
        machine, modules, pm, faults, hmi, screens, vfd,
        physical, lubrication
    """
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    return {
        "machine":     _load_machine(wb),
        "modules":     _load_modules(wb),
        "pm":          _load_pm(wb),
        "faults":      _load_faults(wb),
        "hmi":         _load_hmi(wb),
        "screens":     _load_screens(wb),
        "vfd":         _load_vfd(wb),
        "physical":    _load_sheet_meta(wb, "PHYSICAL"),
        "lubrication": _load_sheet_meta(wb, "LUBRICATION"),
    }


def _load_machine(wb) -> dict:
    """
    SUMMARY sheet rows 2-15: col A = field name, col B = value.
    Returns flat dict of all key→value pairs found.
    """
    ws = wb["SUMMARY"]
    machine: dict[str, str] = {}
    for row in ws.iter_rows(min_row=2, max_row=15, values_only=True):
        key = row[0]
        val = row[1]
        if key and val is not None:
            k = str(key).strip().rstrip(":")
            machine[k] = str(val).strip() if val is not None else ""
    return machine


def _load_modules(wb) -> list[dict]:
    """
    Load every MOD_N sheet as a module with BOM rows.
    Sheet structure:
        Row 4 col B: module display name
        Row 6: column headers (ITEM #, DESCRIPTION, PART #, CONFIG, QTY)
        Row 7+: BOM data
    """
    modules: list[dict] = []
    for sheet_name in wb.sheetnames:
        if not re.match(r"^MOD_\d+$", sheet_name, re.IGNORECASE):
            continue
        ws = wb[sheet_name]

        # Display name is in row 4, col B
        display_name = ws.cell(row=4, column=2).value
        display_name = str(display_name).strip() if display_name else sheet_name

        # BOM: header row 6, data rows 7+
        bom: list[dict] = []
        for row in ws.iter_rows(min_row=7, values_only=True):
            item, desc, part, config, qty = (list(row) + [None] * 5)[:5]
            if item is None and desc is None:
                continue
            # CONFIG cells may be stored as datetime.time (gear-ratio artefact)
            config_str = ""
            if config is not None:
                import datetime
                if isinstance(config, datetime.time):
                    config_str = f"{config.hour}:{config.minute}"
                else:
                    config_str = str(config).strip()
            bom.append({
                "item":        str(item).strip()   if item   is not None else "",
                "description": str(desc).strip()   if desc   is not None else "",
                "part_number": str(part).strip()   if part   is not None else "",
                "config":      config_str,
                "qty":         str(qty).strip()    if qty    is not None else "",
            })

        modules.append({
            "sheet":        sheet_name,
            "display_name": display_name,
            "bom":          bom,
        })

    modules.sort(key=lambda m: int(re.search(r"\d+", m["sheet"]).group()))
    return modules


def _load_pm(wb) -> list[dict]:
    """
    PM sheet: USE=True rows only, cols A-E (USE, ITEM, TASK, FREQUENCY, IMG).
    IMG values that are formula errors are returned as empty string.
    """
    ws = wb["PM"]
    pm: list[dict] = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        use, item, task, freq, img = (list(row) + [None] * 5)[:5]
        # Normalise USE: accept True, "TRUE", "YES", "1"
        if isinstance(use, bool):
            include = use
        elif isinstance(use, str):
            include = use.strip().upper() in ("TRUE", "YES", "1")
        else:
            include = False
        if not include:
            continue

        # Normalise IMG: drop formula-error strings
        img_val = ""
        if img is not None and not str(img).startswith("#"):
            img_val = str(img).strip()

        pm.append({
            "item":      str(item).strip() if item else "",
            "task":      str(task).strip() if task else "",
            "frequency": str(freq).strip() if freq else "",
            "img":       img_val,
        })
    return pm


def _load_faults(wb) -> list[dict]:
    """FAULTS sheet: cols A-C only (NUMBER, NAME, SEVERITY)."""
    ws = wb["FAULTS"]
    faults: list[dict] = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        num, name, severity = (list(row) + [None] * 3)[:3]
        if num is None and name is None:
            continue
        faults.append({
            "number":   str(num).strip()      if num      is not None else "",
            "name":     str(name).strip()     if name     is not None else "",
            "severity": str(severity).strip() if severity is not None else "",
        })
    return faults


def _load_hmi(wb) -> dict:
    """
    HMI sheet: load as a table.
    Row 1: product name headers (col B onward)
    Row 2: product descriptions
    Row 3: OPTION/BUTTON | VALUE headers
    Row 4+: parameter rows
    """
    ws = wb["HMI"]
    all_rows: list[list] = []
    for row in ws.iter_rows(values_only=True):
        cells = [str(c).strip() if c is not None else "" for c in row]
        if any(cells):
            all_rows.append(cells)
    return {"rows": all_rows}


def _load_screens(wb) -> list[dict]:
    """
    SCREENS sheet.
    Row 1: screen names at 1-indexed cols 2, 5, 8, … (every 3rd starting at 2).
    Row 2: "Place in Cell" image per screen (captured via COM).
    Row 3: OPTION / DESCRIPTION column headers.
    Row 4+: per-option data.
        col 3n-2: include checkbox (True/False)
        col 3n-1: option name   (= screen name col)
        col 3n  : description
    where n = screen index (1-based), so screen 1 → cols 1/2/3,
    screen 2 → cols 4/5/6, screen 3 → cols 7/8/9, …
    """
    ws = wb["SCREENS"]

    # Determine screen name columns from row 1
    row1 = [ws.cell(1, c).value for c in range(1, ws.max_column + 1)]
    screen_name_cols: list[tuple[str, int]] = []  # (name, 1-indexed col)
    for col_idx, val in enumerate(row1, start=1):
        if val and isinstance(val, str) and val.strip():
            screen_name_cols.append((val.strip(), col_idx))

    screens: list[dict] = []
    for screen_name, name_col in screen_name_cols:
        checkbox_col = name_col - 1   # 1-indexed
        desc_col     = name_col + 1   # 1-indexed

        options: list[dict] = []
        for row in ws.iter_rows(min_row=SCREENS_DATA_ROW, values_only=True):
            row_list = list(row)
            # 0-indexed access: checkbox @ checkbox_col-1, name @ name_col-1, desc @ desc_col-1
            cb  = row_list[checkbox_col - 1] if checkbox_col >= 1 and len(row_list) >= checkbox_col else None
            opt = row_list[name_col - 1]     if len(row_list) >= name_col     else None
            dsc = row_list[desc_col - 1]     if len(row_list) >= desc_col     else None

            if opt is None and dsc is None:
                continue

            if isinstance(cb, bool):
                include = cb
            elif isinstance(cb, str):
                include = cb.strip().upper() in ("TRUE", "YES", "1")
            else:
                include = False

            options.append({
                "include":     include,
                "option":      str(opt).strip() if opt else "",
                "description": str(dsc).strip() if dsc else "",
            })

        screens.append({
            "name":     screen_name,
            "name_col": name_col,   # used by COM capture to locate row-2 image cell
            "options":  options,
        })

    return screens


def _load_vfd(wb) -> dict:
    """
    Load VFD parameters from the sheet matching the HMI brand in SUMMARY.
    Returns {"brand": ..., "sheet": ..., "rows": [[...], ...]}.
    """
    machine = _load_machine(wb)
    hmi_brand = machine.get("HMI", "").upper()

    sheet_name = VFD_SHEET_MAP.get(hmi_brand)
    if not sheet_name:
        # Partial match fallback
        for key, sname in VFD_SHEET_MAP.items():
            if key in hmi_brand:
                sheet_name = sname
                break

    if not sheet_name or sheet_name not in wb.sheetnames:
        print(f"  WARNING: No VFD sheet found for HMI brand '{hmi_brand}'. "
              f"Available sheets: {wb.sheetnames}")
        return {"brand": hmi_brand, "sheet": None, "rows": []}

    ws = wb[sheet_name]
    rows: list[list] = []
    for row in ws.iter_rows(values_only=True):
        cells = [str(c).strip() if c is not None else "" for c in row]
        if any(cells):
            rows.append(cells)
    return {"brand": hmi_brand, "sheet": sheet_name, "rows": rows}


def _load_sheet_meta(wb, sheet_name: str) -> dict:
    """Return sheet name and openpyxl-reported used range for COM image capture."""
    if sheet_name not in wb.sheetnames:
        return {"sheet": sheet_name, "range": None}
    ws = wb[sheet_name]
    return {"sheet": sheet_name, "range": ws.dimensions}


# ─────────────────────────────────────────────────────────────────────────────
# IMAGE CAPTURE  (Excel COM — Windows + Excel required)
# ─────────────────────────────────────────────────────────────────────────────

def capture_images(xlsx_path: Path, data: dict, out_dir: Path) -> dict[str, Path]:
    """
    Use Excel COM automation to export:
      - PHYSICAL used range     → physical.png
      - LUBRICATION used range  → lubrication.png
      - Each screen's row-2 image cell → screen_00.png, screen_01.png, …

    Returns {key: Path} for each successfully captured image.
    Prints a warning and returns {} if pywin32 or Excel is unavailable.
    """
    if not _WIN32_AVAILABLE:
        print("  WARNING: pywin32 not installed — image capture skipped.\n"
              "           PHYSICAL, LUBRICATION, and SCREENS sections will be omitted.\n"
              "           Run:  pip install pywin32")
        return {}

    out_dir.mkdir(parents=True, exist_ok=True)
    images: dict[str, Path] = {}
    xlsx_abs = str(xlsx_path.resolve())

    xl = win32.Dispatch("Excel.Application")
    xl.Visible = False
    xl.DisplayAlerts = False

    try:
        wb_com = xl.Workbooks.Open(xlsx_abs, ReadOnly=True)

        # ── PHYSICAL and LUBRICATION ──────────────────────────────────────────
        for key in ("physical", "lubrication"):
            meta = data.get(key, {})
            sname = meta.get("sheet")
            rng   = meta.get("range")
            if not sname or not rng:
                continue
            out_png = out_dir / f"{key}.png"
            if _export_range_as_png(wb_com, sname, rng, out_png):
                images[key] = out_png
                print(f"  Captured {key}: {out_png.name}")

        # ── SCREENS row-2 images ──────────────────────────────────────────────
        for i, screen in enumerate(data.get("screens", [])):
            col_letter = get_column_letter(screen["name_col"])
            cell_addr  = f"{col_letter}{SCREENS_IMG_ROW}"
            out_png    = out_dir / f"screen_{i:02d}.png"
            if _export_range_as_png(wb_com, "SCREENS", cell_addr, out_png):
                images[f"screen_{i}"] = out_png
                print(f"  Captured screen {i:02d}: {screen['name']}")
            else:
                print(f"  No image for screen {i:02d}: {screen['name']} (blank or error)")

        wb_com.Close(False)
    except Exception as e:
        print(f"  ERROR during Excel COM image capture: {e}")
    finally:
        try:
            xl.Quit()
        except Exception:
            pass

    return images


def _export_range_as_png(wb_com, sheet_name: str, range_addr: str,
                          out_path: Path) -> bool:
    """
    Export an Excel range (or single cell) as a PNG by pasting it into a
    temporary chart object, then exporting that chart.
    Returns True on success, False on any error.
    """
    try:
        ws = wb_com.Sheets(sheet_name)
        rng = ws.Range(range_addr)

        # CopyPicture: Appearance=1 (xlScreen), Format=2 (xlBitmap)
        rng.CopyPicture(Appearance=1, Format=2)

        w = max(float(rng.Width),  50.0)
        h = max(float(rng.Height), 50.0)

        # Paste into a throw-away chart placed off-screen
        chart_obj = ws.ChartObjects().Add(-2000, -2000, w, h)
        chart     = chart_obj.Chart
        chart.Paste()
        chart.Export(str(out_path.resolve()), "PNG")
        chart_obj.Delete()
        return out_path.exists()
    except Exception as e:
        print(f"    WARNING: capture failed for {sheet_name}!{range_addr}: {e}")
        return False


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT BUILDING
# ─────────────────────────────────────────────────────────────────────────────

def build_document(data: dict, images: dict[str, Path],
                   template_path: Path, out_path: Path):
    """
    Assemble the manual Word document from data + captured images.

    Sections inserted (in order):
        Cover placeholders  ← text substitution in template
        Modules BOM         ← one table per MOD_N
        PM                  ← filtered table
        Faults              ← table
        HMI Params          ← table
        HMI Screens         ← image + option table per screen
        Setpoints           ← single image (PHYSICAL)
        Lubrication         ← single image (LUBRICATION)
        VFD Params          ← table
    """
    machine = data["machine"]

    out_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(template_path, out_path)
    doc = Document(str(out_path))

    # ── Cover substitutions ───────────────────────────────────────────────────
    serial   = machine.get("SERIAL NUMBER", "")
    customer = machine.get("CUSTOMER",      "CUSTOMER")
    mtype    = machine.get("TYPE",          "MODEL")
    date     = machine.get("DELIVERY DATE", machine.get("MANUAL DATE", ""))

    # Year (YY) for filename — strip century if full year given
    year_raw = machine.get("YEAR (YY)", "")
    if not year_raw and date:
        m = re.search(r"\b(\d{4})\b", date)
        if m:
            year_raw = m.group(1)[-2:]

    # SN short (last 4 digits)
    sn_digits = re.sub(r"\D", "", serial)
    sn_short  = sn_digits[-4:] if len(sn_digits) >= 4 else sn_digits

    replacements = {
        "CUSTOMER":         customer,
        "MODEL: MODEL":     f"MODEL: {mtype}",
        "SERIAL: 2X-4XXX":  f"SERIAL: {serial}",
        "MMMM-YYYY":        date,
    }
    for old, new in replacements.items():
        _replace_text(doc, old, new)

    # ── Modules ───────────────────────────────────────────────────────────────
    for mod in data.get("modules", []):
        doc.add_heading(mod["display_name"], level=1)
        if mod["bom"]:
            _insert_bom_table(doc, mod["bom"])
        else:
            doc.add_paragraph("(No BOM data)", style="Normal")

    # ── Preventive Maintenance ────────────────────────────────────────────────
    pm_rows = data.get("pm", [])
    if pm_rows:
        doc.add_heading("Preventive Maintenance", level=1)
        _insert_pm_table(doc, pm_rows)

    # ── Faults ────────────────────────────────────────────────────────────────
    faults = data.get("faults", [])
    if faults:
        doc.add_heading("Fault Codes", level=1)
        _insert_faults_table(doc, faults)

    # ── HMI Parameters ───────────────────────────────────────────────────────
    hmi = data.get("hmi", {})
    if hmi.get("rows"):
        doc.add_heading("HMI Parameters", level=1)
        _insert_generic_table(doc, hmi["rows"])

    # ── HMI Screens ──────────────────────────────────────────────────────────
    screens = data.get("screens", [])
    if screens:
        doc.add_heading("HMI Screens", level=1)
        for i, screen in enumerate(screens):
            doc.add_heading(screen["name"], level=2)
            img_key = f"screen_{i}"
            if img_key in images:
                _insert_centered_image(doc, images[img_key], width=DInches(5.5))
            options = [o for o in screen["options"] if o["include"]]
            if options:
                _insert_screen_options_table(doc, options)

    # ── Physical Setpoints ────────────────────────────────────────────────────
    if "physical" in images:
        doc.add_heading("Physical Setpoints", level=1)
        _insert_centered_image(doc, images["physical"], width=DInches(6.0))

    # ── Lubrication ──────────────────────────────────────────────────────────
    if "lubrication" in images:
        doc.add_heading("Lubrication", level=1)
        _insert_centered_image(doc, images["lubrication"], width=DInches(6.0))

    # ── VFD Parameters ───────────────────────────────────────────────────────
    vfd = data.get("vfd", {})
    if vfd.get("rows"):
        brand = vfd.get("brand", "VFD")
        doc.add_heading(f"VFD Parameters — {brand.title()}", level=1)
        _insert_generic_table(doc, vfd["rows"])

    doc.save(str(out_path))
    print(f"\n  Saved: {out_path}\n")


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _replace_text(doc: Document, old: str, new: str):
    """Replace all occurrences of old→new in paragraphs and table cells."""
    for para in doc.paragraphs:
        if old in para.text:
            for run in para.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in para.text:
                        for run in para.runs:
                            if old in run.text:
                                run.text = run.text.replace(old, new)


def _insert_bom_table(doc: Document, bom: list[dict]):
    headers = ["ITEM #", "DESCRIPTION", "PART #", "CONFIG", "QTY"]
    table = doc.add_table(rows=1 + len(bom), cols=len(headers))
    table.style = "Table Grid"
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for ri, row_data in enumerate(bom):
        vals = [row_data["item"], row_data["description"], row_data["part_number"],
                row_data["config"], row_data["qty"]]
        for ci, v in enumerate(vals):
            table.rows[ri + 1].cells[ci].text = v
    doc.add_paragraph()


def _insert_pm_table(doc: Document, pm: list[dict]):
    headers = ["ITEM", "TASK", "FREQUENCY", "IMG"]
    table = doc.add_table(rows=1 + len(pm), cols=len(headers))
    table.style = "Table Grid"
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for ri, row_data in enumerate(pm):
        vals = [row_data["item"], row_data["task"],
                row_data["frequency"], row_data["img"]]
        for ci, v in enumerate(vals):
            table.rows[ri + 1].cells[ci].text = v
    doc.add_paragraph()


def _insert_faults_table(doc: Document, faults: list[dict]):
    headers = ["ALARM #", "FAULT NAME", "SEVERITY"]
    table = doc.add_table(rows=1 + len(faults), cols=len(headers))
    table.style = "Table Grid"
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for ri, f in enumerate(faults):
        vals = [f["number"], f["name"], f["severity"]]
        for ci, v in enumerate(vals):
            table.rows[ri + 1].cells[ci].text = v
    doc.add_paragraph()


def _insert_generic_table(doc: Document, rows: list[list]):
    """Insert a plain table from a list-of-lists. First row is bold header."""
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for ri, row_data in enumerate(rows):
        for ci in range(max_cols):
            val = row_data[ci] if ci < len(row_data) else ""
            cell = table.rows[ri].cells[ci]
            cell.text = val
            if ri == 0:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
    doc.add_paragraph()


def _insert_screen_options_table(doc: Document, options: list[dict]):
    headers = ["OPTION", "DESCRIPTION"]
    table = doc.add_table(rows=1 + len(options), cols=2)
    table.style = "Table Grid"
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for ri, opt in enumerate(options):
        table.rows[ri + 1].cells[0].text = opt["option"]
        table.rows[ri + 1].cells[1].text = opt["description"]
    doc.add_paragraph()


def _insert_centered_image(doc: Document, img_path: Path, width=None):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    kwargs = {}
    if width:
        kwargs["width"] = width
    run.add_picture(str(img_path), **kwargs)
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# MAIN PIPELINE
# ─────────────────────────────────────────────────────────────────────────────

def generate(config_path: Path, json_only: bool = False):
    if not json_only:
        print(f"\n  Config : {config_path}")

    # 1. Load all data
    data = load_data(config_path)
    machine = data["machine"]

    # ── --json mode: dump data and exit ──────────────────────────────────────
    if json_only:
        print(json.dumps(data, indent=2, default=str))
        return

    # 2. Determine output filename
    serial   = machine.get("SERIAL NUMBER", "0000")
    customer = machine.get("CUSTOMER",      "CUSTOMER").upper()
    mtype    = machine.get("TYPE",          "MODEL").upper()
    date     = machine.get("DELIVERY DATE", machine.get("MANUAL DATE", ""))

    sn_digits = re.sub(r"\D", "", serial)
    sn_short  = sn_digits[-4:] if len(sn_digits) >= 4 else sn_digits

    year_raw = machine.get("YEAR (YY)", "")
    if not year_raw and date:
        m = re.search(r"\b(\d{4})\b", date)
        if m:
            year_raw = m.group(1)[-2:]
    yy = year_raw or "XX"

    # Sanitise customer/type for filename
    safe_customer = re.sub(r'[\\/*?:"<>|]', "", customer)
    safe_type     = re.sub(r'[\\/*?:"<>|]', "", mtype)
    out_name = f"{yy}-{sn_short} - {safe_customer} - {safe_type}.docx"
    out_path = OUTPUT_DIR / out_name

    print(f"  Output : {out_path}")
    print(f"  Modules: {[m['display_name'] for m in data['modules']]}")
    print(f"  VFD    : {data['vfd'].get('sheet', 'none')}")
    print(f"  PM rows: {len(data['pm'])}")
    print(f"  Faults : {len(data['faults'])}")
    print(f"  Screens: {len(data['screens'])}\n")

    if not TEMPLATE_DOCX.exists():
        sys.exit(f"ERROR: {TEMPLATE_DOCX} not found.")

    # 3. Capture images via Excel COM
    with tempfile.TemporaryDirectory() as tmp_root:
        tmp_dir = Path(tmp_root)
        images  = capture_images(config_path, data, tmp_dir)

        # 4. Build the document
        build_document(data, images, TEMPLATE_DOCX, out_path)


# ─────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Polypack Manual Generator v2")
    parser.add_argument(
        "--config", "-c",
        required=True,
        type=Path,
        help="Path to the machine Excel workbook (.xlsx)"
    )
    parser.add_argument(
        "--json",
        action="store_true",
        help="Dump extracted data as JSON and exit (no Word doc generated)"
    )
    args = parser.parse_args()
    generate(args.config, json_only=args.json)
